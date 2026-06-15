"""Extract plain text from HTML, RSS, or search snippets."""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Any, Optional

import feedparser
import trafilatura
from bs4 import BeautifulSoup

try:  # gói mới (đã đổi tên); fallback gói cũ nếu chưa cài
    from ddgs import DDGS
except ImportError:  # pragma: no cover
    from duckduckgo_search import DDGS


def normalize_for_hash(text: str) -> str:
    t = text.strip().lower()
    t = re.sub(r"\s+", " ", t)
    return t


def is_pdf_bytes(data: bytes) -> bool:
    """Magic bytes for PDF (works after optional BOM/whitespace)."""
    if not data:
        return False
    return data.lstrip()[:4] == b"%PDF"


def extract_from_pdf_bytes(data: bytes) -> str:
    """
    Extract plain text from a PDF with a text layer.
    Scanned-only PDFs (image) return a short placeholder.
    """
    if not is_pdf_bytes(data):
        return ""
    try:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                parts.append(t.strip())
        out = "\n\n".join(parts).strip()
        if out:
            return out
        return (
            "(PDF: không có lớp chữ — có thể là file scan/ảnh; "
            "cần OCR nếu muốn trích nội dung.)"
        )
    except Exception as exc:  # noqa: BLE001 — surface as text for pipeline
        return f"(PDF: không đọc được — {type(exc).__name__}: {exc})"


def extract_from_docx_bytes(data: bytes) -> str:
    """Extract plain text from a .docx file."""
    try:
        from io import BytesIO

        from docx import Document

        doc = Document(BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip() or "(DOCX: file rỗng hoặc không có nội dung text.)"
    except Exception as exc:  # noqa: BLE001
        return f"(DOCX: không đọc được — {type(exc).__name__}: {exc})"


def extract_from_html(html: str, *, mode: str) -> str:
    if mode == "raw":
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    # favor_recall + include_links: giữ thêm khối nội dung phụ / liên kết so với mặc định
    # (vẫn có thể mất sidebar so với raw — dùng extract=raw nếu cần đủ mọi chữ trên trang).
    extracted = trafilatura.extract(
        html,
        include_comments=False,
        include_tables=True,
        no_fallback=False,
        favor_recall=True,
        favor_precision=False,
        include_links=True,
        include_formatting=False,
    )
    if extracted and extracted.strip():
        return extracted.strip()
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n", strip=True)


def _rss_field_to_plain(html_or_text: str) -> str:
    s = (html_or_text or "").strip()
    if not s:
        return ""
    if "<" in s and ">" in s:
        soup = BeautifulSoup(s, "html.parser")
        return soup.get_text(separator="\n", strip=True)
    return s


def extract_from_rss_feed(xml_bytes: bytes, *, max_entries: int) -> str:
    parsed = feedparser.parse(xml_bytes)
    lines: list[str] = []
    for entry in parsed.entries[:max_entries]:
        title = entry.get("title", "")
        link = entry.get("link", "")
        summary = entry.get("summary", entry.get("description", ""))
        body_html = summary
        raw_content = entry.get("content")
        if isinstance(raw_content, list) and raw_content:
            merged: list[str] = []
            for block in raw_content:
                if not isinstance(block, dict):
                    continue
                val = (block.get("value") or "").strip()
                if val:
                    merged.append(val)
            joined = "\n\n".join(merged).strip()
            if len(joined) > len(summary or ""):
                body_html = joined
        body = _rss_field_to_plain(str(body_html))
        lines.append(f"# {title}\n{link}\n{body}\n")
    return "\n".join(lines).strip() or "(empty feed)"


def search_to_text(query: str, *, max_results: int) -> str:
    import warnings

    lines: list[str] = []
    results: list[dict[str, str]] = []
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        try:
            with DDGS() as ddgs:
                # Thử thêm region — một số môi trường Bing/DDG trả rỗng với mặc định.
                for reg in (None, "us-en", "wt-wt"):
                    kwargs: dict[str, object] = {"max_results": max_results}
                    if reg is not None:
                        kwargs["region"] = reg
                    results = list(ddgs.text(query, **kwargs))
                    if results:
                        break
        except Exception as exc:  # noqa: BLE001
            return (
                f"(search error: {type(exc).__name__}: {exc})\n"
                "Gợi ý: thử lại sau, đổi mạng/VPN, hoặc dùng URL/RSS trực tiếp thay vì Search."
            )
    for r in results:
        title = r.get("title", "")
        href = r.get("href", "")
        body = r.get("body", "")
        lines.append(f"{title}\n{href}\n{body}\n")
    if not lines:
        return (
            "(no search results — DuckDuckGo/Bing trả về rỗng trong môi trường này. "
            "Dùng nguồn URL hoặc RSS cụ thể, hoặc cập nhật gói duckduckgo-search / thử mạng khác.)"
        )
    return "\n".join(lines).strip()


def image_search_results(query: str, *, max_results: int = 200, regions=("wt-wt", "us-en", "vn-vi")) -> list[tuple[str, str]]:
    """Tìm ảnh qua ddgs.images → trả list (url, title), gộp nhiều region, bỏ trùng url."""
    import warnings

    seen: set[str] = set()
    out: list[tuple[str, str]] = []
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        for reg in regions:
            try:
                with DDGS() as ddgs:
                    for r in ddgs.images(query, region=reg, max_results=max_results):
                        u = (r.get("image") or "").strip()
                        if u.startswith("http") and u not in seen:
                            seen.add(u)
                            out.append((u, str(r.get("title") or "")))
            except Exception:
                continue
            if len(out) >= max_results:
                break
    return out


def image_search_urls(query: str, *, max_results: int = 200, regions=("wt-wt", "us-en", "vn-vi")) -> list[str]:
    """Tìm ảnh qua ddgs.images → trả list URL ảnh trực tiếp (gộp nhiều region, bỏ trùng)."""
    return [u for u, _ in image_search_results(query, max_results=max_results, regions=regions)]


@dataclass
class SourcePayload:
    text: str
    meta: dict[str, Any]
    etag: Optional[str] = None
    last_modified: Optional[str] = None
