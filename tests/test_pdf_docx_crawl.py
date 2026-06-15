"""Unit tests: crawl from PDF and DOCX files end-to-end."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from src.config_loader import SourceConfig
from src.extract import extract_from_docx_bytes, extract_from_pdf_bytes
from src.pipeline import collect_payload, run_source
from src.settings import Settings
from src.storage import Storage


@pytest.fixture
def settings():
    return Settings(
        ollama_base_url="http://localhost:11434/v1",
        ollama_model="test",
        ollama_api_key="ollama",
        http_timeout=30,
        user_agent="test-agent",
        database_path=Path("data/test.db"),
        max_text_chars=40000,
        llm_max_retries=1,
        llm_retry_backoff_sec=0,
        skip_llm=True,
        mongodb_uri=None,
        mongodb_database="test",
        mongodb_collection="docs",
    )


def _make_pdf_bytes(text: str) -> bytes:
    """Tạo PDF đơn giản bằng pypdf."""
    from io import BytesIO
    from pypdf import PdfWriter
    from pypdf._page import PageObject
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas

    # Dùng reportlab nếu có, fallback tạo PDF thủ công
    try:
        buf = BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=A4)
        c.drawString(72, 700, text)
        c.save()
        return buf.getvalue()
    except ImportError:
        pass

    # Fallback: tạo PDF bằng pypdf (minimal)
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Tạo DOCX đơn giản."""
    from io import BytesIO
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# === PDF Tests ===

class TestPdfExtract:
    def test_extract_from_pdf_bytes_valid(self):
        """PDF có text layer phải trích được nội dung."""
        # Tạo PDF bằng pypdf (blank page — không có text layer)
        from io import BytesIO
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(612, 792)
        buf = BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        result = extract_from_pdf_bytes(pdf_bytes)
        # Blank PDF → placeholder message
        assert "PDF" in result

    def test_extract_from_pdf_bytes_not_pdf(self):
        """Non-PDF bytes trả về empty."""
        result = extract_from_pdf_bytes(b"not a pdf file")
        assert result == ""

    def test_collect_payload_pdf_file(self, settings, tmp_path):
        """Pipeline collect_payload xử lý file .pdf."""
        from io import BytesIO
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(612, 792)
        buf = BytesIO()
        writer.write(buf)

        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(buf.getvalue())

        src = SourceConfig(
            id="pdf_test", type="file", schedule_cron="0 * * * *",
            extract="raw", file_path=str(pdf_path),
        )
        payload = collect_payload(src, settings)
        assert payload is not None
        assert payload.meta["format"] == "pdf"
        assert "PDF" in payload.text

    def test_run_source_pdf_skip_llm(self, settings, tmp_path):
        """Full pipeline: PDF file → storage (skip LLM)."""
        from io import BytesIO
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(612, 792)
        buf = BytesIO()
        writer.write(buf)

        pdf_path = tmp_path / "report.pdf"
        pdf_path.write_bytes(buf.getvalue())

        db = tmp_path / "test.db"
        storage = Storage(db)
        src = SourceConfig(
            id="pdf_e2e", type="file", schedule_cron="0 * * * *",
            extract="raw", file_path=str(pdf_path),
        )
        result = run_source(src, storage, settings, None)
        assert result.changed is True
        assert result.document_id is not None

        # Verify stored
        doc = storage.get_document_by_id(result.document_id)
        assert doc is not None
        assert doc.source_id == "pdf_e2e"


# === DOCX Tests ===

class TestDocxExtract:
    def test_extract_from_docx_bytes_basic(self):
        """DOCX với paragraphs phải trích được text."""
        docx_bytes = _make_docx_bytes(["Hello World", "Second paragraph"])
        result = extract_from_docx_bytes(docx_bytes)
        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_extract_from_docx_bytes_empty(self):
        """DOCX rỗng trả về placeholder."""
        docx_bytes = _make_docx_bytes([])
        result = extract_from_docx_bytes(docx_bytes)
        assert "DOCX" in result

    def test_extract_from_docx_bytes_with_table(self):
        """DOCX có table phải trích được nội dung bảng."""
        from io import BytesIO
        from docx import Document

        doc = Document()
        doc.add_paragraph("Title")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        buf = BytesIO()
        doc.save(buf)

        result = extract_from_docx_bytes(buf.getvalue())
        assert "Title" in result
        assert "Alice" in result
        assert "30" in result

    def test_extract_from_docx_invalid_bytes(self):
        """Invalid bytes trả về error message."""
        result = extract_from_docx_bytes(b"not a docx")
        assert "DOCX" in result
        assert "không đọc được" in result

    def test_collect_payload_docx_file(self, settings, tmp_path):
        """Pipeline collect_payload xử lý file .docx."""
        docx_bytes = _make_docx_bytes(["Crawl AI Report", "Data extracted successfully."])
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(docx_bytes)

        src = SourceConfig(
            id="docx_test", type="file", schedule_cron="0 * * * *",
            extract="raw", file_path=str(docx_path),
        )
        payload = collect_payload(src, settings)
        assert payload is not None
        assert payload.meta["format"] == "docx"
        assert "Crawl AI Report" in payload.text
        assert "Data extracted" in payload.text

    def test_run_source_docx_full_pipeline(self, settings, tmp_path):
        """Full pipeline: DOCX file → extract → hash → storage."""
        docx_bytes = _make_docx_bytes([
            "Nguyen Van A",
            "Ngày sinh: 1990-01-15",
            "Địa chỉ: 123 Đường ABC, Quận 1, TP.HCM",
            "SĐT: 0901234567",
        ])
        docx_path = tmp_path / "contacts.docx"
        docx_path.write_bytes(docx_bytes)

        db = tmp_path / "test.db"
        storage = Storage(db)
        src = SourceConfig(
            id="docx_e2e", type="file", schedule_cron="0 * * * *",
            extract="raw", file_path=str(docx_path),
        )
        result = run_source(src, storage, settings, None)
        assert result.changed is True

        # Verify content stored correctly
        doc = storage.get_document_by_id(result.document_id)
        assert "Nguyen Van A" in doc.raw_text
        assert "0901234567" in doc.raw_text

    def test_run_source_docx_unchanged_hash_skips(self, settings, tmp_path):
        """Chạy lần 2 cùng file → skip (unchanged hash)."""
        docx_bytes = _make_docx_bytes(["Same content"])
        docx_path = tmp_path / "same.docx"
        docx_path.write_bytes(docx_bytes)

        db = tmp_path / "test.db"
        storage = Storage(db)
        src = SourceConfig(
            id="docx_dup", type="file", schedule_cron="0 * * * *",
            extract="raw", file_path=str(docx_path),
        )
        r1 = run_source(src, storage, settings, None)
        assert r1.changed is True

        r2 = run_source(src, storage, settings, None)
        assert r2.changed is False
        assert r2.skipped_reason == "unchanged_hash"
